"""
파일 처리 모듈
다양한 파일 형식을 처리하고 텍스트로 추출
"""
import os
import io
import json
import base64
from PIL import Image
try:
    from pypdf import PdfReader  # PyPDF2 3.0+ (pypdf로 변경됨)
except ImportError:
    try:
        from PyPDF2 import PdfReader  # PyPDF2 구버전 호환성
    except ImportError:
        PdfReader = None  # PDF 처리 불가능
from docx import Document
import openpyxl

# Pandas는 선택적 임포트 (Vercel 서버리스 크기 제한 대응)
# Vercel 환경에서는 pandas를 사용하지 않고 openpyxl만 사용
import os
IS_VERCEL = os.environ.get('IS_VERCEL', 'false').lower() == 'true'

if not IS_VERCEL:
    try:
        import pandas as pd
        PANDAS_AVAILABLE = True
    except ImportError:
        pd = None
        PANDAS_AVAILABLE = False
else:
    # Vercel에서는 pandas 사용 안 함
    pd = None
    PANDAS_AVAILABLE = False


class FileProcessor:
    """파일 처리 클래스"""

    # 지원하는 파일 확장자
    SUPPORTED_EXTENSIONS = {
        'image': ['.jpg', '.jpeg', '.png', '.gif', '.webp'],
        'document': ['.pdf', '.docx', '.txt'],
        'spreadsheet': ['.xlsx', '.csv'],
        'data': ['.json', '.xml']
    }

    # 최대 파일 크기 (100MB)
    MAX_FILE_SIZE = 100 * 1024 * 1024

    @staticmethod
    def is_supported(filename):
        """파일이 지원되는 형식인지 확인"""
        ext = os.path.splitext(filename)[1].lower()
        for extensions in FileProcessor.SUPPORTED_EXTENSIONS.values():
            if ext in extensions:
                return True
        return False

    @staticmethod
    def get_file_type(filename):
        """파일 타입 반환"""
        ext = os.path.splitext(filename)[1].lower()
        for file_type, extensions in FileProcessor.SUPPORTED_EXTENSIONS.items():
            if ext in extensions:
                return file_type
        return 'unknown'

    @staticmethod
    def process_image(file_data, filename):
        """이미지 파일 처리"""
        try:
            image = Image.open(io.BytesIO(file_data))

            # 이미지 정보
            info = {
                'type': 'image',
                'filename': filename,
                'size': f"{image.size[0]}x{image.size[1]}",
                'format': image.format,
                'mode': image.mode
            }

            # base64 인코딩 (GPT-4 Vision용)
            base64_image = base64.b64encode(file_data).decode('utf-8')
            info['base64'] = base64_image

            return {
                'success': True,
                'info': info,
                'text': f"[이미지: {filename}, 크기: {info['size']}, 형식: {image.format}]",
                'has_image': True
            }
        except Exception as e:
            return {
                'success': False,
                'error': f"이미지 처리 오류: {str(e)}"
            }

    @staticmethod
    def process_pdf(file_data, filename):
        """PDF 파일 처리"""
        if PdfReader is None:
            return {
                'success': False,
                'error': 'PDF 처리 라이브러리가 설치되지 않았습니다. pypdf 또는 PyPDF2를 설치해주세요.'
            }

        try:
            pdf_file = io.BytesIO(file_data)
            reader = PdfReader(pdf_file)

            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"

            return {
                'success': True,
                'info': {
                    'type': 'pdf',
                    'filename': filename,
                    'pages': len(reader.pages)
                },
                'text': text.strip() if text.strip() else "[PDF에서 텍스트를 추출할 수 없습니다]",
                'has_image': False
            }
        except Exception as e:
            import traceback
            print(f"[PDF 처리 오류] {filename}: {str(e)}")
            traceback.print_exc()
            return {
                'success': False,
                'error': f"PDF 처리 오류: {str(e)}"
            }

    @staticmethod
    def process_docx(file_data, filename):
        """DOCX 파일 처리"""
        try:
            docx_file = io.BytesIO(file_data)
            doc = Document(docx_file)

            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            return {
                'success': True,
                'info': {
                    'type': 'docx',
                    'filename': filename,
                    'paragraphs': len(doc.paragraphs)
                },
                'text': text.strip() if text.strip() else "[DOCX에서 텍스트를 추출할 수 없습니다]",
                'has_image': False
            }
        except Exception as e:
            return {
                'success': False,
                'error': f"DOCX 처리 오류: {str(e)}"
            }

    @staticmethod
    def process_txt(file_data, filename):
        """TXT 파일 처리"""
        try:
            text = file_data.decode('utf-8')

            return {
                'success': True,
                'info': {
                    'type': 'txt',
                    'filename': filename,
                    'lines': len(text.split('\n'))
                },
                'text': text,
                'has_image': False
            }
        except Exception as e:
            return {
                'success': False,
                'error': f"TXT 처리 오류: {str(e)}"
            }

    @staticmethod
    def process_xlsx(file_data, filename):
        """XLSX 파일 처리"""
        try:
            print(f"[FileProcessor] XLSX 처리 시작: {filename}, 크기: {len(file_data)} bytes")
            xlsx_file = io.BytesIO(file_data)

            # Pandas가 사용 가능하면 pandas 사용 (빠르고 편리함)
            if PANDAS_AVAILABLE:
                # pandas로 Excel 읽기 - 모든 시트 읽기
                all_sheets = pd.read_excel(xlsx_file, engine='openpyxl', sheet_name=None)
                print(f"[FileProcessor] Excel 읽기 성공 (pandas): {len(all_sheets)}개 시트")

                # 모든 시트의 데이터 수집
                all_table_data = []
                text_parts = []
                total_rows = 0
                total_columns = 0

                for sheet_name, df in all_sheets.items():
                    print(f"[FileProcessor] 시트 '{sheet_name}': {len(df)} 행, {len(df.columns)} 열")

                    # 데이터프레임을 텍스트로 변환
                    sheet_text = f"\n[시트: {sheet_name}]\n"
                    sheet_text += f"행: {len(df)}, 열: {len(df.columns)}\n"
                    sheet_text += f"열 이름: {', '.join(df.columns.tolist())}\n\n"
                    sheet_text += df.to_string()
                    text_parts.append(sheet_text)

                    # 표 데이터 추출 (전체 행)
                    preview_df = df

                    # NaN 값을 빈 문자열로 변환
                    sheet_table_data = {
                        'sheet_name': sheet_name,
                        'headers': preview_df.columns.tolist(),
                        'rows': preview_df.fillna('').values.tolist(),
                        'total_rows': len(df),
                        'total_columns': len(df.columns)
                    }
                    all_table_data.append(sheet_table_data)

                    total_rows += len(df)
                    total_columns = max(total_columns, len(df.columns))

                # 전체 텍스트 결합
                text = "\n\n".join(text_parts)
                summary = f"총 {len(all_sheets)}개 시트, 전체 행: {total_rows}\n\n"

                print(f"[FileProcessor] XLSX 처리 완료 (pandas): {filename}, {len(all_sheets)}개 시트")
                return {
                    'success': True,
                    'info': {
                        'type': 'xlsx',
                        'filename': filename,
                        'rows': total_rows,
                        'columns': total_columns,
                        'sheets': len(all_sheets)
                    },
                    'text': summary + text,
                    'has_image': False,
                    'table_data': all_table_data[0] if all_table_data else None,
                    'all_sheets': all_table_data
                }

            else:
                # Pandas 없으면 openpyxl만 사용 (Vercel 환경)
                print(f"[FileProcessor] openpyxl만 사용하여 처리 (pandas 없음)")
                wb = openpyxl.load_workbook(xlsx_file, data_only=True)

                all_table_data = []
                text_parts = []
                total_rows = 0
                total_columns = 0

                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]

                    # 헤더와 데이터 읽기
                    rows_data = []
                    headers = []

                    for row_idx, row in enumerate(ws.iter_rows(values_only=True), 1):
                        if row_idx == 1:
                            # 첫 번째 행을 헤더로 사용
                            headers = [str(cell) if cell is not None else f'Column{i+1}' for i, cell in enumerate(row)]
                        else:
                            # 데이터 행
                            row_data = [cell if cell is not None else '' for cell in row]
                            rows_data.append(row_data)

                    if not headers:
                        continue

                    num_rows = len(rows_data)
                    num_cols = len(headers)

                    print(f"[FileProcessor] 시트 '{sheet_name}': {num_rows} 행, {num_cols} 열")

                    # 텍스트 변환
                    sheet_text = f"\n[시트: {sheet_name}]\n"
                    sheet_text += f"행: {num_rows}, 열: {num_cols}\n"
                    sheet_text += f"열 이름: {', '.join(headers)}\n\n"

                    # 표 형식으로 데이터 추가 (최대 100행만)
                    for i, row in enumerate(rows_data[:100]):
                        sheet_text += f"  {', '.join(str(cell) for cell in row)}\n"
                    if num_rows > 100:
                        sheet_text += f"  ... ({num_rows - 100}개 행 더 있음)\n"

                    text_parts.append(sheet_text)

                    # 표 데이터 저장
                    sheet_table_data = {
                        'sheet_name': sheet_name,
                        'headers': headers,
                        'rows': rows_data,
                        'total_rows': num_rows,
                        'total_columns': num_cols
                    }
                    all_table_data.append(sheet_table_data)

                    total_rows += num_rows
                    total_columns = max(total_columns, num_cols)

                # 전체 텍스트 결합
                text = "\n\n".join(text_parts)
                summary = f"총 {len(wb.sheetnames)}개 시트, 전체 행: {total_rows}\n\n"

                print(f"[FileProcessor] XLSX 처리 완료 (openpyxl): {filename}, {len(wb.sheetnames)}개 시트")
                return {
                    'success': True,
                    'info': {
                        'type': 'xlsx',
                        'filename': filename,
                        'rows': total_rows,
                        'columns': total_columns,
                        'sheets': len(wb.sheetnames)
                    },
                    'text': summary + text,
                    'has_image': False,
                    'table_data': all_table_data[0] if all_table_data else None,
                    'all_sheets': all_table_data
                }

        except Exception as e:
            print(f"[FileProcessor] XLSX 처리 오류: {filename} - {str(e)}")
            import traceback
            traceback.print_exc()
            return {
                'success': False,
                'error': f"XLSX 처리 오류: {str(e)}"
            }

    @staticmethod
    def process_csv(file_data, filename):
        """CSV 파일 처리"""
        try:
            if PANDAS_AVAILABLE:
                # pandas 사용
                csv_file = io.BytesIO(file_data)
                df = pd.read_csv(csv_file)

                text = df.to_string()
                summary = f"행: {len(df)}, 열: {len(df.columns)}\n열 이름: {', '.join(df.columns.tolist())}\n\n"

                # 표 데이터 추출 (전체 행)
                preview_df = df
                table_data = {
                    'headers': preview_df.columns.tolist(),
                    'rows': preview_df.values.tolist(),
                    'total_rows': len(df),
                    'total_columns': len(df.columns)
                }

                return {
                    'success': True,
                    'info': {
                        'type': 'csv',
                        'filename': filename,
                        'rows': len(df),
                        'columns': len(df.columns)
                    },
                    'text': summary + text,
                    'has_image': False,
                    'table_data': table_data
                }
            else:
                # pandas 없이 csv 모듈 사용
                import csv
                text_data = file_data.decode('utf-8')
                csv_reader = csv.reader(io.StringIO(text_data))

                rows = list(csv_reader)
                if not rows:
                    return {
                        'success': False,
                        'error': 'CSV 파일이 비어있습니다.'
                    }

                headers = rows[0]
                data_rows = rows[1:]

                # 텍스트 변환
                text = f"행: {len(data_rows)}, 열: {len(headers)}\n"
                text += f"열 이름: {', '.join(headers)}\n\n"

                # 데이터 추가 (최대 100행)
                for i, row in enumerate(data_rows[:100]):
                    text += f"  {', '.join(str(cell) for cell in row)}\n"
                if len(data_rows) > 100:
                    text += f"  ... ({len(data_rows) - 100}개 행 더 있음)\n"

                table_data = {
                    'headers': headers,
                    'rows': data_rows,
                    'total_rows': len(data_rows),
                    'total_columns': len(headers)
                }

                return {
                    'success': True,
                    'info': {
                        'type': 'csv',
                        'filename': filename,
                        'rows': len(data_rows),
                        'columns': len(headers)
                    },
                    'text': text,
                    'has_image': False,
                    'table_data': table_data
                }

        except Exception as e:
            return {
                'success': False,
                'error': f"CSV 처리 오류: {str(e)}"
            }

    @staticmethod
    def process_json(file_data, filename):
        """JSON 파일 처리"""
        try:
            text = file_data.decode('utf-8')
            data = json.loads(text)

            # JSON을 보기 좋게 포맷팅
            formatted = json.dumps(data, indent=2, ensure_ascii=False)

            return {
                'success': True,
                'info': {
                    'type': 'json',
                    'filename': filename
                },
                'text': formatted,
                'has_image': False
            }
        except Exception as e:
            return {
                'success': False,
                'error': f"JSON 처리 오류: {str(e)}"
            }

    @staticmethod
    def process_xml(file_data, filename):
        """XML 파일 처리"""
        try:
            text = file_data.decode('utf-8')

            return {
                'success': True,
                'info': {
                    'type': 'xml',
                    'filename': filename
                },
                'text': text,
                'has_image': False
            }
        except Exception as e:
            return {
                'success': False,
                'error': f"XML 처리 오류: {str(e)}"
            }

    @staticmethod
    def process_file(file_data, filename):
        """파일 타입에 따라 적절한 처리 함수 호출"""
        ext = os.path.splitext(filename)[1].lower()

        # 파일 크기 체크
        if len(file_data) > FileProcessor.MAX_FILE_SIZE:
            return {
                'success': False,
                'error': '파일 크기가 너무 큽니다 (최대 100MB)'
            }

        # 확장자별 처리
        if ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp']:
            return FileProcessor.process_image(file_data, filename)
        elif ext == '.pdf':
            return FileProcessor.process_pdf(file_data, filename)
        elif ext == '.docx':
            return FileProcessor.process_docx(file_data, filename)
        elif ext == '.txt':
            return FileProcessor.process_txt(file_data, filename)
        elif ext == '.xlsx':
            return FileProcessor.process_xlsx(file_data, filename)
        elif ext == '.csv':
            return FileProcessor.process_csv(file_data, filename)
        elif ext == '.json':
            return FileProcessor.process_json(file_data, filename)
        elif ext == '.xml':
            return FileProcessor.process_xml(file_data, filename)
        else:
            return {
                'success': False,
                'error': f'지원하지 않는 파일 형식입니다: {ext}'
            }
